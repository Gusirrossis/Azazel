"""Plugins de documentos: PDF (pypdf) y DOCX (python-docx). Python puro — sin JVM.

PDF: primero texto NATIVO (pypdf). Si el PDF es un ESCANEO (texto nativo casi vacío) y
`ctx.ocr_activo`, se rasterizan sus páginas (pypdfium2) y se les hace OCR (Fase 2). Todo
degrada con flags: sin pypdfium2/tesseract el PDF se indexa igual con lo que haya.

El bucle de OCR respeta el PLAZO cooperativo (`ctx.vencido()`): al agotarse devuelve
las páginas que alcanzó con la bandera `ocr_pdf_parcial`. Antes, el corte por timeout
ocurría fuera y descartaba el resultado entero — un PDF de 20 páginas que llegaba a
la 15 se indexaba sin una línea de texto.
"""

from __future__ import annotations

from typing import Any

from . import ContextoExtraccion, ResultadoExtraccion, registrar
from ._ocr import ocr_imagen


def _ocr_pdf(ctx: ContextoExtraccion) -> tuple[str | None, list[str], float | None]:
    """Rasteriza páginas y las pasa por OCR. Solo para escaneos. Nunca lanza."""
    try:
        import pypdfium2 as pdfium  # dependencia opcional (extra `ocr`)
    except Exception:
        return None, ["ocr_no_disponible"], None
    p = ctx.perillas
    ctx.fuente.seek(0)
    try:
        pdf = pdfium.PdfDocument(ctx.fuente.read())
    except Exception as exc:  # PDF cifrado/corrupto → flag, no rompe
        return None, [f"ocr_pdf_fallido:{type(exc).__name__}"], None

    textos: list[str] = []
    flags: set[str] = set()
    # Confianza ponderada por longitud: una página de 900 caracteres pesa más que
    # una carátula de 12, igual que dentro de una página pesan más las palabras largas.
    suma_conf = 0.0
    suma_peso = 0.0
    total = len(pdf)
    limite = min(total, p.ocr_pdf_max_paginas)
    paginas_hechas = 0
    try:
        for i in range(limite):
            if ctx.vencido():
                flags.add("ocr_pdf_parcial")
                break
            pagina = pdf[i]
            imagen = pagina.render(scale=p.ocr_pdf_escala).to_pil()
            texto_pag, flags_pag, conf_pag = ocr_imagen(imagen, p)
            paginas_hechas += 1
            if texto_pag:
                textos.append(texto_pag)
                if conf_pag is not None:
                    peso = float(len(texto_pag))
                    suma_conf += conf_pag * peso
                    suma_peso += peso
            flags.update(f for f in flags_pag if f != "ocr_ok")
            if sum(len(t) for t in textos) >= p.extractor_max_chars:
                flags.add("texto_truncado")
                break
    finally:
        pdf.close()

    if total > limite:
        flags.add("ocr_pdf_paginas_limitadas")
    texto = "\n".join(textos)[: p.extractor_max_chars].strip()
    confianza = round(suma_conf / suma_peso, 1) if suma_peso > 0 else None
    if not texto:
        return None, sorted(flags) or ["ocr_vacio"], None
    return texto, ["ocr_pdf", f"ocr_pdf_paginas:{paginas_hechas}", *sorted(flags)], confianza


@registrar("application/pdf")
def extraer_pdf(ctx: ContextoExtraccion) -> ResultadoExtraccion:
    from pypdf import PdfReader

    lector = PdfReader(ctx.fuente)
    campos: dict[str, Any] = {"paginas": len(lector.pages)}
    meta = lector.metadata
    if meta is not None:
        if meta.title:
            campos["titulo"] = str(meta.title)[:300]
        if meta.author:
            campos["autor"] = str(meta.author)[:300]

    maximo = ctx.perillas.extractor_max_chars
    fragmentos: list[str] = []
    acumulado = 0
    flags: list[str] = []
    for pagina in lector.pages:
        # También aquí: un PDF con miles de páginas de texto nativo puede agotar el
        # plazo antes del OCR, y lo leído hasta ahora vale.
        if ctx.vencido():
            flags.append("texto_parcial")
            break
        fragmento = pagina.extract_text() or ""
        fragmentos.append(fragmento)
        acumulado += len(fragmento)
        if acumulado >= maximo:
            flags.append("texto_truncado")  # patrón fscrawler: indexed_chars (⚙K11)
            break
    texto = "\n".join(fragmentos)[:maximo].strip()

    # Escaneo (texto nativo casi vacío) + OCR activo → rasterizar y OCR (Fase 2).
    if ctx.ocr_activo and len(texto) < ctx.perillas.ocr_pdf_umbral_chars and not ctx.vencido():
        texto_ocr, flags_ocr, confianza = _ocr_pdf(ctx)
        if texto_ocr:
            return ResultadoExtraccion(
                campos=campos, texto=texto_ocr, flags=flags_ocr, confianza=confianza
            )
        flags = [*flags, *flags_ocr]

    return ResultadoExtraccion(campos=campos, texto=texto or None, flags=flags)


@registrar("application/vnd.openxmlformats-officedocument.wordprocessingml.document")
def extraer_docx(ctx: ContextoExtraccion) -> ResultadoExtraccion:
    from docx import Document

    documento = Document(ctx.fuente)
    parrafos = [p.text for p in documento.paragraphs if p.text.strip()]
    campos: dict[str, Any] = {"parrafos": len(parrafos)}
    propiedades = documento.core_properties
    if propiedades.title:
        campos["titulo"] = propiedades.title[:300]
    if propiedades.author:
        campos["autor"] = propiedades.author[:300]

    maximo = ctx.perillas.extractor_max_chars
    texto = "\n".join(parrafos)
    flags = ["texto_truncado"] if len(texto) > maximo else []
    return ResultadoExtraccion(campos=campos, texto=texto[:maximo] or None, flags=flags)
